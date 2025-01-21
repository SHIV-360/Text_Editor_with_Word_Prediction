import tkinter as tk
from tkinter import font
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt

suggest = "example"  # Word suggestion variable

def toggle_formatting(tag, font_style=None):
    """Toggle formatting like bold, italic, or underline on selected text."""
    try:
        start_index = text.index("sel.first")
        end_index = text.index("sel.last")

        if tag in text.tag_names("sel.first"):
            text.tag_remove(tag, "sel.first", "sel.last")
        else:
            text.tag_add(tag, "sel.first", "sel.last")
            text.tag_configure(tag, font=(current_font.get(), 12, font_style))
    except tk.TclError:
        messagebox.showwarning("Warning", "No text selected to apply formatting!")

def change_font(event=None):
    """Change font style for the selected text."""
    selected_font = current_font.get()
    text.config(font=(selected_font, 12))

def toggle_case(case_type):
    """Toggle between uppercase and lowercase for selected text."""
    try:
        start_index = text.index("sel.first")
        end_index = text.index("sel.last")
        selected_text = text.get("sel.first", "sel.last")
        new_text = (
            selected_text.upper() if case_type == "uppercase" else selected_text.lower()
        )
        text.delete("sel.first", "sel.last")
        text.insert(start_index, new_text)
    except tk.TclError:
        messagebox.showwarning("Warning", "No text selected to toggle case!")

def save_as_docx():
    """Save the content to a DOCX file with proper formatting."""
    file_name = file_name_entry.get()
    if not file_name:
        messagebox.showwarning("Warning", "Please enter a file name before saving!")
        return

    try:
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", initialfile=file_name, filetypes=[("Word Files", "*.docx")]
        )
        if file_path:
            doc = Document()

            # Get the selected font from the current font variable
            selected_font = current_font.get()

            # Add the text content to the DOCX file with formatting
            for line_index, line in enumerate(text.get("1.0", "end").splitlines()):
                paragraph = doc.add_paragraph()

                tags = text.tag_names(f"{line_index + 1}.0")
                run = paragraph.add_run(line)

                # Apply font styling based on selected tags
                if "bold" in tags:
                    run.bold = True
                if "italic" in tags:
                    run.italic = True
                if "underline" in tags:
                    run.underline = True

                run.font.name = selected_font
                run.font.size = Pt(12)  # Set the font size to 12pt

            doc.save(file_path)
            messagebox.showinfo("Success", "File saved successfully as DOCX!")
    except Exception as e:
        messagebox.showerror("Error", f"Could not save file: {e}")

def toggle_theme():
    """Toggle between light and dark mode."""
    if theme_button.config('text')[-1] == 'Light Mode':
        text.config(bg="white", fg="black", insertbackground="black")
        root.config(bg="white")
        toolbar.config(bg="white")
        for widget in toolbar.winfo_children():
            widget.config(bg="white", fg="black")
        theme_button.config(text="Dark Mode")
    else:
        text.config(bg="black", fg="white", insertbackground="white")
        root.config(bg="black")
        toolbar.config(bg="black")
        for widget in toolbar.winfo_children():
            widget.config(bg="black", fg="white")
        theme_button.config(text="Light Mode")

def update_last_words(event=None):
    """Fetch and display the last five words in the terminal."""
    current_text = text.get("1.0", "end").strip()
    words = current_text.split()
    last_five_words = words[-5:] if len(words) > 5 else words
    print("Last 5 words:", last_five_words)

def show_suggestion(event=None):
    """Display the word suggestion in the GUI."""
    suggestion_label.config(text=f"Suggested word: {suggest}")

def insert_suggestion(event=None):
    """Insert the suggested word at the current cursor position."""
    text.insert(tk.INSERT, f"{suggest} ")

# Create the main application window
root = tk.Tk()
root.title("Enhanced Text Editor")
root.geometry("800x600")

current_font = tk.StringVar(value="Arial")

toolbar = tk.Frame(root)
toolbar.pack(fill="x", padx=5, pady=5)

font_label = tk.Label(toolbar, text="Font:")
font_label.pack(side="left", padx=5)

font_menu = tk.OptionMenu(toolbar, current_font, *font.families(), command=change_font)
font_menu.pack(side="left", padx=5)

bold_btn = tk.Button(toolbar, text="Bold", command=lambda: toggle_formatting("bold", "bold"))
bold_btn.pack(side="left", padx=5)

italic_btn = tk.Button(toolbar, text="Italic", command=lambda: toggle_formatting("italic", "italic"))
italic_btn.pack(side="left", padx=5)

underline_btn = tk.Button(toolbar, text="Underline", command=lambda: toggle_formatting("underline", "underline"))
underline_btn.pack(side="left", padx=5)

uppercase_btn = tk.Button(toolbar, text="Toggle Case", command=lambda: toggle_case("uppercase"))
uppercase_btn.pack(side="left", padx=5)

file_name_label = tk.Label(toolbar, text="File Name:")
file_name_label.pack(side="right", padx=5)

file_name_entry = tk.Entry(toolbar, width=20)
file_name_entry.pack(side="right", padx=5)

save_docx_btn = tk.Button(toolbar, text="Save as DOCX", command=save_as_docx)
save_docx_btn.pack(side="right", padx=5)

theme_button = tk.Button(toolbar, text="Light Mode", command=toggle_theme)
theme_button.pack(side="right", padx=5)

# Suggestion Label
suggestion_label = tk.Label(root, text=f"Suggested word: {suggest}", font=("Arial", 10), fg="blue")
suggestion_label.pack(fill="x")

# Text Area
text = tk.Text(root, wrap="word", font=("Arial", 12))
text.pack(expand=1, fill="both", padx=10, pady=10)

# Bind events
text.bind("<KeyRelease>", update_last_words)
text.bind("<FocusIn>", show_suggestion)
text.bind("<Tab>", insert_suggestion)

root.mainloop()
