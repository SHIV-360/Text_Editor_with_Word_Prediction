import tkinter as tk
from tkinter import font
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, RGBColor

def toggle_formatting(tag, font_style=None):
    """Toggle formatting like bold, italic, or underline on selected text."""
    try:
        start_index = text.index("sel.first")
        end_index = text.index("sel.last")

        # Check if the tag is already applied and toggle
        if tag in text.tag_names("sel.first"):
            text.tag_remove(tag, "sel.first", "sel.last")
        else:
            text.tag_add(tag, "sel.first", "sel.last")
            # Apply the correct font style
            text.tag_configure(tag, font=(current_font.get(), 12, font_style))
    except tk.TclError:
        messagebox.showwarning("Warning", "No text selected to apply formatting!")

def change_font(event=None):
    """Change font style for the selected text."""
    selected_font = current_font.get()
    text.config(font=(selected_font, 12))

def toggle_case(case_type):
    """Toggle between uppercase and lowercase for selected text."""
    try:
        start_index = text.index("sel.first")
        end_index = text.index("sel.last")
        selected_text = text.get("sel.first", "sel.last")
        new_text = (
            selected_text.upper() if case_type == "uppercase" else selected_text.lower()
        )
        if selected_text == new_text:
            new_text = selected_text.lower() if case_type == "uppercase" else selected_text.upper()

        text.delete("sel.first", "sel.last")
        text.insert(start_index, new_text)
    except tk.TclError:
        messagebox.showwarning("Warning", "No text selected to toggle case!")

def save_as_docx():
    """Save the content to a DOCX file with proper formatting."""
    file_name = file_name_entry.get()
    if not file_name:
        messagebox.showwarning("Warning", "Please enter a file name before saving!")
        return

    try:
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", initialfile=file_name, filetypes=[("Word Files", "*.docx")]
        )
        if file_path:
            doc = Document()

            # Get the selected font from the current font variable
            selected_font = current_font.get()

            # Add the text content to the DOCX file with formatting
            for line_index, line in enumerate(text.get("1.0", "end").splitlines()):
                paragraph = doc.add_paragraph()

                tags = text.tag_names(f"{line_index + 1}.0")
                run = paragraph.add_run(line)

                # Apply font styling based on selected tags
                if "bold" in tags:
                    run.bold = True
                if "italic" in tags:
                    run.italic = True
                if "underline" in tags:
                    run.underline = True

                run.font.name = selected_font
                run.font.size = Pt(12)  # Set the font size to 12pt

            doc.save(file_path)
            messagebox.showinfo("Success", "File saved successfully as DOCX!")
    except Exception as e:
        messagebox.showerror("Error", f"Could not save file: {e}")

def toggle_theme():
    """Toggle between light and dark mode."""
    if theme_button.config('text')[-1] == 'Light Mode':
        text.config(bg="white", fg="black", insertbackground="black")
        root.config(bg="white")
        toolbar.config(bg="white")
        for widget in toolbar.winfo_children():
            widget.config(bg="white", fg="black")
        theme_button.config(text="Dark Mode")
    else:
        text.config(bg="black", fg="white", insertbackground="white")
        root.config(bg="black")
        toolbar.config(bg="black")
        for widget in toolbar.winfo_children():
            widget.config(bg="black", fg="white")
        theme_button.config(text="Light Mode")

# Create the main application window
root = tk.Tk()
root.title("Enhanced Text Editor")
root.geometry("800x600")

# Initialize the default theme as Dark Mode
root.config(bg="black")

available_fonts = font.families()
current_font = tk.StringVar(value="Arial")

toolbar = tk.Frame(root, bg="black")
toolbar.pack(fill="x", padx=5, pady=5)

font_label = tk.Label(toolbar, text="Font:", bg="black", fg="white")
font_label.pack(side="left", padx=5)

# Updated the font menu to show a sample text in the selected font
font_menu = tk.OptionMenu(
    toolbar, current_font, *available_fonts, command=change_font
)
font_menu.config(bg="black", fg="white")
font_menu["menu"].config(bg="black", fg="white")
font_menu.pack(side="left", padx=5)

bold_btn = tk.Button(toolbar, text="Bold", command=lambda: toggle_formatting("bold", "bold"), bg="black", fg="white")
bold_btn.pack(side="left", padx=5)

italic_btn = tk.Button(toolbar, text="Italic", command=lambda: toggle_formatting("italic", "italic"), bg="black", fg="white")
italic_btn.pack(side="left", padx=5)

underline_btn = tk.Button(toolbar, text="Underline", command=lambda: toggle_formatting("underline", "underline"), bg="black", fg="white")
underline_btn.pack(side="left", padx=5)

uppercase_btn = tk.Button(toolbar, text="Toggle Case", command=lambda: toggle_case("uppercase"), bg="black", fg="white")
uppercase_btn.pack(side="left", padx=5)

file_name_label = tk.Label(toolbar, text="File Name:", bg="black", fg="white")
file_name_label.pack(side="right", padx=5)
file_name_entry = tk.Entry(toolbar, width=20, bg="black", fg="white", insertbackground="white")
file_name_entry.pack(side="right", padx=5)

save_docx_btn = tk.Button(toolbar, text="Save as DOCX", command=save_as_docx, bg="black", fg="white")
save_docx_btn.pack(side="right", padx=5)

theme_button = tk.Button(toolbar, text="Light Mode", command=toggle_theme, bg="black", fg="white")
theme_button.pack(side="right", padx=5)

text = tk.Text(root, wrap="word", font=("Arial", 12), bg="black", fg="white", insertbackground="white")
text.pack(expand=1, fill="both", padx=10, pady=10)

root.mainloop()
